# #import 
# import streamlit as st
# import pandas as pd
# import numpy as np
# from wordcloud import wordcloud, STOPWORDS
# import PyPDF2
# from docx import Document
# import base64
# from io import BytesIO
# import plotly.express as px
# import os

# #functions for text file reading
# def read_text(file):
#     return file.getvalue().decode( "utf-8" )
# #functions for docx file reading
# def read_docx(file):
#     doc = Document(file)
#     return "".join([para.text for para in doc.paragraphs])

# #functions for pdf file reading
# # Extract text from PDF
# def read_pdf(file):
#     pdfReader = PyPDF2.PdfFileReader(file)
#     return "".join([page.extractText() for page in pdfReader.pages])    
   
# #functions to filter out stop words
# def remove_stopwords(text):
#     stop_words = set(STOPWORDS)
#     return " ".join([word for word in str(text).split() if word not in stop_words])
# #functions to create  download link for plot
# def create_download_link(figure, title):
#     file = BytesIO()
#     figure.savefig(file, format="png")
#     file.seek(0)
#     b64 = base64.b64encode(file.read()).decode()
#     return f'<a href="data:image/png;base64,{b64}" download="{title}.png">Download {title} plot</a>'

# #function to create a download link for a dataframe
# def create_download_link_df(df, title):#this function will create a download link for a dataframe
#     csv = df.to_csv(index=False)#convert the dataframe to csv
#     b64 = base64.b64encode(csv.encode()).decode()# some strings <-> bytes conversions necessary here 
#     return f'<a href="data:file/csv;base64,{b64}" download="{title}.csv">Download {title} file</a>'#define the filename here 

# #title
# st.title("Welcome to  Word Cloud App!")
# st.subheader("Upload your file here")

# uploaded_file = st.file_uploader("Choose a file", type= ("pdf", "docx", "txt"))#upload the file
# st.set_option('deprecation.showfileUploaderEncoding', False)#to remove the warning message

# #set condition for file type

# if uploaded_file:
#     if uploaded_file.type == "application/pdf":
#         text = read_pdf(uploaded_file)
#         st.write(text)
#     elif uploaded_file.type == "text/plain":
#         text = read_text(uploaded_file)
#         st.write(text)
#     elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         text = read_docx(uploaded_file)
#         st.write(text)
#     else:
#         st.error("File format not supported")
#         st.stop()

#     # Generate word count table
#     words = text.split()  # Use 'text' instead of 'text'
#     word_count = pd.DataFrame({'word': words}).groupby('word').size().reset_index(name='Count').sort_values(by='Count', ascending=False)

#     # Sidebar checkbox and multi-select box for stopwords
#     use_standard_stopwords = st.sidebar.checkbox("Use standard stopwords")
#     top_words=word_count.head(50).tolist()#get the top 50 words from the word count table
#     additional_stopwords = st.sidebar.multiselect("Additional stopwords", top_words, default="")

#     if use_standard_stopwords:
#         all_stopwords = STOPWORDS.union(additional_stopwords)#union the standard stopwords and additional stopwords to get all the stopwords
#     else:
#         all_stopwords = set(additional_stopwords)#set the additional stopwords as all the stopwords
         
#     text = filter(remove_stopwords, text, all_stopwords)

# if text:
#     #word cloud dimension
#     width = st.sidebar.slider("Width", 200, 400)
#     height = st.sidebar.slider("Height", 200, 400)

#     #create word cloud
#     import matplotlib.pyplot as plt
# st.sidebar.subheader("Generate Word Cloud") 
# fig, ax = plt.subplots(figsize=(width/100, height/100))
# wordcloud_img=wordcloud.WordCloud(width=width, height=height, background_color="white", max_words=100).generate(text)
# ax.imshow(wordcloud_img, interpolation='bilinear')
# ax.axis("off")

# #save the plot functionality
# format=st.sidebar.selectbox("Select the format", ["png", "pdf", "svg"])
# resolution=st.sidebar.slider("Resolution", 100, 300)

# #generate word count table
# st.subheader("Word Count Table")
# words=text.split()
# word_count = pd.DataFrame({'word': words}).groupby('word').size().reset_index(name='Count').sort_values(by='Count', ascending=False)
# st.write(word_count)
# st.pyplot(fig)

# if st.button("save as {format}"):
#     Buffer = BytesIO()
#     plt.savefig(Buffer, format=format, dpi=resolution)
#     st.markdown(create_download_link(fig, "wordcloud"), unsafe_allow_html=True)


#     #word count table at the end
#     st.sidebar.subheader("the creater of the Word Cloud is Muhammad Asif")
    
import streamlit as st
import pandas as pd
from wordcloud import WordCloud, STOPWORDS
import base64
from io import BytesIO
import matplotlib.pyplot as plt
from PyPDF2 import PdfReader
from docx import Document

# Functions for file reading
def read_text(file):
    return file.getvalue().decode("utf-8")

def read_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def read_pdf(file):
    pdf_reader = PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

# Function to remove stopwords
def remove_stopwords(text, stopwords):
    return " ".join([word for word in str(text).split() if word not in stopwords])

# Function to create download link for plot
def create_download_link(figure, title, file_format):
    if file_format == "png":
        file_extension = "png"
    elif file_format == "pdf":
        file_extension = "pdf"
    elif file_format == "svg":
        file_extension = "svg"
    else:
        raise ValueError("Invalid file format specified.")
    
    buffer = BytesIO()
    figure.savefig(buffer, format=file_extension)
    buffer.seek(0)
    b64 = base64.b64encode(buffer.read()).decode()
    return f'<a href="data:image/{file_extension};base64,{b64}" download="{title}.{file_extension}">Download {title} plot</a>'

# Streamlit app
def main():
    st.title("Word Cloud App")
    st.sidebar.title("Settings")

    uploaded_file = st.sidebar.file_uploader("Upload File", type=["txt", "pdf", "docx"])

    if uploaded_file:
        if uploaded_file.type == "text/plain":
            text = read_text(uploaded_file)
        elif uploaded_file.type == "application/pdf":
            text = read_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = read_docx(uploaded_file)
        else:
            st.error("Unsupported file format")
            st.stop()

        # Additional options
        st.sidebar.subheader("Word Cloud Options")
        background_color = st.sidebar.color_picker("Background Color", "#FFFFFF")
        max_words = st.sidebar.slider("Max Words", min_value=50, max_value=500, value=200)
        width = st.sidebar.slider("Width", min_value=200, max_value=800, value=400)
        height = st.sidebar.slider("Height", min_value=200, max_value=800, value=400)
        stopwords = st.sidebar.text_area("Stopwords (comma-separated)", "")

        # Word cloud generation
        st.subheader("Word Cloud")
        stopwords = set(STOPWORDS.union(stopwords.split(",")))
        text_no_stopwords = remove_stopwords(text, stopwords)
        wordcloud = WordCloud(stopwords=stopwords, background_color=background_color, max_words=max_words, width=width, height=height).generate(text_no_stopwords)
        plt.figure(figsize=(10, 6))
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis("off")
        st.pyplot(plt)

        # Save word cloud
        st.subheader("Save Word Cloud")
        file_format = st.selectbox("Select file format", ["png", "pdf", "svg"])
        if st.button(f"Save Word Cloud as {file_format.upper()}"):
            download_link = create_download_link(plt, "wordcloud", file_format)
            st.markdown(download_link, unsafe_allow_html=True)

if __name__ == "__main__":
    main()

